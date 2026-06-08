"""DOCX export service."""

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from models.schemas import DraftContent


def generate_docx(final_draft: DraftContent, document_type: str) -> io.BytesIO:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    normal = doc.styles["Normal"].font
    normal.name = "Times New Roman"
    normal.size = Pt(12)

    title = doc.add_heading(final_draft.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    for section_data in final_draft.sections:
        doc.add_heading(section_data.heading, level=2)
        paragraph = doc.add_paragraph(section_data.content)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
